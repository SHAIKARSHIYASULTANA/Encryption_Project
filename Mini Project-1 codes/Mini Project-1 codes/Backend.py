import os
import docx
from datetime import datetime
import pytesseract
import io
import PyPDF2
from io import BytesIO
import base64
from PIL import Image
from PyPDF2 import PdfFileReader
import fitz 
from flask import Flask, render_template, request, Response,url_for

app = Flask(__name__)

def encrypt(text):
    # Encrypt the text using Caesar Cipher
    t = ''
    n1=1 
    n2=1
    for i in text:
        np=0
        ntg=0
        np=int((n1*(n1+1)*(n1+2))/6)
        x=ord(i)
        y=(x^np)
        ntg=int((n2*(n2+1)*(n2+2)*(n2+3))/24)
        f=y//ntg
        r=y%ntg
        t=t+'f'+str(f)+'r'+str(r)
        n1=n1+1
        n2=n2+1
    return t

def decrypt(text):
    # Encrypt the text using Caesar Cipher
    t=text
    ori = ''
    n1=1 
    n2=1
    i=0
    while(i!=len(t)):
        z=''
        rem=''
        dividend=0
        if(t[i]=='f'):
            i=i+1
            while(t[i]!='r'):
                z=z+t[i]
                i=i+1
            if(t[i]=='r'):
                i=i+1
                while(i!=len(t) and t[i]!='f'):
                    rem=rem+t[i]
                    i=i+1
        ntr=int((n2*(n2+1)*(n2+2)*(n2+3))/24)
        np=int((n1*(n1+1)*(n1+2))/6)
        m=ntr*int(z)
        dividend=m+int(rem)
        msg=dividend^np
        ori=ori+chr(msg)
        n1=n1+1 
        n2=n2+1

    return ori


def read_pdf(file_data):
    pdf = fitz.open(stream=file_data, filetype="pdf")
    text = ""
    for page in pdf:
        text += page.getText()
    return text

@app.route('/')
def index():
    return render_template('index.html')
@app.route('/result', methods=['POST'])
def result():
    current_time_str =datetime.now().strftime("%H:%M:%S")
    uploaded_time = datetime.now()
    timestamp_upload = datetime.now().strftime("%H:%M:%S")
    mode = request.form['mode']
    if mode == 'encrypt':
        text_file = request.files['text_file']
        original_name = text_file.filename
        original_size=round(float(len(text_file.read())) / 1024, 2)
        text_file.seek(0)

        filename = text_file.filename
        extension = os.path.splitext(filename)[1]

        result_text=''
        # Read the contents of the text file
        if extension == '.txt':
            text = text_file.read().decode('utf-8')
            result_text = encrypt(text)
            text_file.seek(0)
        elif extension == '.doc' or extension == '.docx':
            document = docx.Document(text_file)
            text = '\n'.join([paragraph.text for paragraph in document.paragraphs])
            result_text = encrypt(text)
            text_file.seek(0)
        else:
            return "Error: File type not supported"

        with open(filename, 'w') as f:
            f.write(result_text)

        file_type = 'encrypted'
        size = os.path.getsize(filename) / 1024
        download_link = f'<a href="/download/{filename}">Download Encrypted Text ({size:.2f} KB)</a>'
        uploaded_time_str = uploaded_time.strftime('%H:%M:%S')
        uploaded_time_diff = datetime.now() - uploaded_time
        uploaded_time_diff_str = str(uploaded_time_diff)
        file_path = os.path.dirname(os.path.abspath(text_file.filename))
    else:
        encrypted_file = request.files['encrypted_file']
        original_name = encrypted_file.filename
        original_size=round(float(len(encrypted_file.read())) / 1024, 2)
        encrypted_file.seek(0)
        uploaded_time = datetime.now()

        filename = encrypted_file.filename
        extension = os.path.splitext(filename)[1]
        decrypted_text=''
        # Read the contents of the encrypted file
        if extension == '.txt':
            encrypted_text = encrypted_file.read().decode('utf-8')
            decrypted_text = decrypt(encrypted_text)
            encrypted_file.seek(0)
        elif extension == '.doc' or extension == '.docx':
            document = docx.Document(encrypted_file)
            encrypted_text = '\n'.join([paragraph.text for paragraph in document.paragraphs])
            decrypted_text = decrypt(encrypted_text)
            encrypted_file.seek(0)
        else:
            return "Error: Invalid file type. Please select an encrypted text file."


        with open(filename, 'w') as f:
            f.write(decrypted_text)

        file_type = 'decrypted'
        size = os.path.getsize(filename) / 1024
        download_link = f'<a href="/download/{file_type}">Download Decrypted Text ({size:.2f} KB)</a>'
         # set the upload time
        uploaded_time_str = uploaded_time.strftime('%H:%M:%S')
        uploaded_time_diff = datetime.now() - uploaded_time
        uploaded_time_diff_str = str(uploaded_time_diff)
        file_path = os.path.dirname(os.path.abspath(encrypted_file.filename))
    return render_template('index.html', original_name=original_name, original_size=original_size, uploaded_time=uploaded_time_str,filename=filename, download_link=download_link, file_size=size, file_type=file_type,current_time=current_time_str,uploaded_time_diff=uploaded_time_diff_str,file_path=file_path)


@app.route('/download/<file_type>/<path:filepath>')
def download(file_type, filepath):
    
    filename = os.path.basename(filepath)

    if os.path.exists(os.path.abspath(filepath)):
        with open(filepath, 'rb') as f:
            file_content = f.read()
        if file_type == 'encrypted':
            return Response(
                file_content,
                mimetype="text/plain",
                headers={"Content-disposition":
                         f"attachment; filename=encrypted_{filename}"})
        else:
            return Response(
                file_content,
                mimetype="text/plain",
                headers={"Content-disposition":
                         f"attachment; filename=decrypted_{filename}"})
    else:
        return "Error: File not found"





if __name__ == '__main__':
    app.run(debug=True)
